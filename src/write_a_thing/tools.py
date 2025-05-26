"""The tools used by the agents."""

import os

from bs4 import BeautifulSoup, Tag
from docling.document_converter import DocumentConverter
from docling.exceptions import ConversionError
from docx import Document as DocxDocument
from markdown import markdown
from smolagents import tool


@tool
def ask_user(question: str) -> str:
    """Ask the user a question and return their response.

    Args:
        question:
            The question to ask the user.

    Returns:
        The user's response to the question.
    """
    return input(f"Question for you: {question}\n> ")


@tool
def load_document(file_path: str) -> str:
    """Load a document from the given file path.

    The `file_path` should point to an existing document file.

    Args:
        file_path:
            The path to the document file.

    Returns:
        The Markdown parsed content of the document.
    """
    try:
        converter = DocumentConverter()
        docling_doc = converter.convert(source=file_path).document
        return docling_doc.export_to_markdown()
    except ConversionError:
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()


@tool
def save_as_word(markdown_content: str, output_path: str) -> None:
    """Save the given Markdown content as a Word document.

    Args:
        markdown_content:
            The Markdown content to save as a Word document.
        output_path:
            The path where the Word document will be saved.

    Raises:
        FileExistsError: If the output file already exists.
        ValueError: If the content could not be parsed.
    """
    if os.path.exists(output_path):
        raise FileExistsError(
            f"The file {output_path} already exists. Please choose a different name."
        )

    doc = DocxDocument()
    html = markdown(text=markdown_content, extensions=["tables"])
    soup = BeautifulSoup(html, "html.parser")
    for element in soup.contents:
        assert isinstance(element, Tag), "The `element` is not a Tag!"
        if element.name is None:
            raise ValueError(f"Invalid HTML content provided: {html}")
        if element.name and element.name.startswith("h"):
            level = int(element.name[1])
            doc.add_heading(element.get_text(), level=level)
        elif element.name == "p":
            doc.add_paragraph(element.get_text())
        elif element.name in ["ul", "ol"]:
            style = "List Bullet" if element.name == "ul" else "List Number"
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style=style)
        elif element.name == "blockquote":
            doc.add_paragraph(element.get_text(), style="Intense Quote")
        elif element.name == "table":
            rows = element.find_all("tr")
            if rows:
                first_row = rows[0]
                assert isinstance(first_row, Tag), "The `first_row` is not a Tag!"
                cols = first_row.find_all(["td", "th"])
                table = doc.add_table(rows=len(rows), cols=len(cols))
                for i, row in enumerate(rows):
                    assert isinstance(row, Tag), "The `row` is not a Tag!"
                    for j, cell in enumerate(row.find_all(["td", "th"])):
                        table.cell(i, j).text = cell.get_text()
    doc.save(output_path)
